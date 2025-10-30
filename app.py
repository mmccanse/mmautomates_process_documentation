import streamlit as st
import google.generativeai as genai
import os
from dotenv import load_dotenv
import tempfile
import json
from datetime import datetime
from moviepy.editor import VideoFileClip
import time
import cv2
from PIL import Image
import io
import base64
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pickle
import re

# Load environment variables
load_dotenv()

# Configure Gemini API
if 'GEMINI_API_KEY' in os.environ:
    genai.configure(api_key=os.environ['GEMINI_API_KEY'])
    model = genai.GenerativeModel('gemini-2.5-pro')
else:
    model = None

# Page config
st.set_page_config(
    page_title="AI Process Documentation Generator",
    page_icon="📹",
    layout="wide"
)

# Custom CSS
st.markdown("""
<style>
    .main-header {
        text-align: center;
        padding: 2rem 0;
        background: linear-gradient(90deg, #667eea 0%, #764ba2 100%);
        color: white;
        border-radius: 10px;
        margin-bottom: 2rem;
    }
    .timestamp-item {
        background: #f0f2f6;
        padding: 10px;
        margin: 5px 0;
        border-radius: 5px;
        border-left: 4px solid #667eea;
    }
    .success-box {
        background: #d4edda;
        border: 1px solid #c3e6cb;
        border-radius: 5px;
        padding: 15px;
        margin: 10px 0;
    }
    .moment-card {
        background: white;
        border: 1px solid #e0e0e0;
        border-radius: 8px;
        padding: 15px;
        margin: 10px 0;
        box-shadow: 0 2px 4px rgba(0,0,0,0.1);
    }
    .image-viewer-container {
        position: sticky;
        top: 0;
        z-index: 100;
        background: white;
        padding: 20px;
        border-radius: 10px;
        box-shadow: 0 4px 6px rgba(0,0,0,0.1);
        margin-bottom: 20px;
    }
    /* Hide Streamlit's image expand icon */
    button[title="View fullscreen"] {
        display: none !important;
    }
    [data-testid="stImage"] button {
        display: none !important;
    }
    .image-viewer-container img {
        margin: 0 auto;
        display: block;
    }

    /* Hide Streamlit's inline uploaded filename row under the uploader */
    [data-testid="stFileUploader"] [data-testid="stUploadedFile"],
    [data-testid="stFileUploader"] .uploadedFile {
        display: none !important;
    }

</style>
""", unsafe_allow_html=True)

def initialize_session_state():
    """Initialize session state variables"""
    if 'transcript' not in st.session_state:
        st.session_state.transcript = None
    if 'key_moments' not in st.session_state:
        st.session_state.key_moments = None
    if 'video_path' not in st.session_state:
        st.session_state.video_path = None
    if 'audio_path' not in st.session_state:
        st.session_state.audio_path = None
    if 'extracted_frames' not in st.session_state:
        st.session_state.extracted_frames = None
    if 'word_doc_bytes' not in st.session_state:
        st.session_state.word_doc_bytes = None
    if 'editing_mode' not in st.session_state:
        st.session_state.editing_mode = False
    if 'google_creds' not in st.session_state:
        st.session_state.google_creds = None
    if 'moments_to_delete' not in st.session_state:
        st.session_state.moments_to_delete = set()
    if 'viewing_image' not in st.session_state:
        st.session_state.viewing_image = None

def extract_audio_from_video(video_path):
    """Extract audio from video file"""
    try:
        with st.spinner("📹 Extracting audio from video..."):
            # Verify file exists and has content
            if not os.path.exists(video_path):
                raise FileNotFoundError(f"Video file not found: {video_path}")
            
            file_size = os.path.getsize(video_path)
            if file_size == 0:
                raise ValueError("Video file is empty")
            
            st.info(f"Video file size: {file_size / (1024*1024):.2f} MB")
            
            # Try to open video with error handling
            try:
                video = VideoFileClip(video_path)
            except Exception as e:
                st.error("The video file appears to be corrupted or incompatible.")
                st.error("Please try:")
                st.error("1. Re-exporting the video in MP4 format")
                st.error("2. Using a smaller file size")
                st.error("3. Using a different recording tool")
                raise e
            
            # Check if video has audio
            if video.audio is None:
                video.close()
                raise ValueError("Video file has no audio track. Please record with audio enabled.")
            
            # Create temporary audio file
            audio_path = tempfile.NamedTemporaryFile(delete=False, suffix='.mp3').name
            video.audio.write_audiofile(audio_path, verbose=False, logger=None)
            
            # Get video duration
            duration = video.duration
            
            video.close()
            
            return audio_path, duration
            
    except Exception as e:
        st.error(f"Error extracting audio: {str(e)}")
        return None, None

def transcribe_audio_with_gemini(audio_path):
    """Transcribe audio using Gemini API"""
    try:
        with st.spinner("🎤 Transcribing audio with AI... This may take a minute..."):
            # Read audio file
            with open(audio_path, 'rb') as f:
                audio_bytes = f.read()
            
            # Create prompt
            prompt = """
            Please transcribe this audio recording. 
            
            This is someone explaining an accounting or business process.
            
            Provide the transcript with approximate timestamps in the format [MM:SS] at the beginning of each major sentence or thought.
            
            Be accurate with the transcription, including any technical terms, system names, or navigation paths mentioned.
            """
            
            # Generate transcription using File API
            import mimetypes
            mime_type = mimetypes.guess_type(audio_path)[0] or 'audio/mpeg'
            
            # Create file part
            audio_part = {
                "mime_type": mime_type,
                "data": audio_bytes
            }
            
            # Generate content
            response = model.generate_content([prompt, audio_part])
            
            return response.text
            
    except Exception as e:
        st.error(f"Error transcribing audio: {str(e)}")
        import traceback
        st.error(f"Details: {traceback.format_exc()}")
        return None

def analyze_transcript_for_key_moments(transcript):
    """Analyze transcript to identify key moments for screenshots"""
    try:
        with st.spinner("🧠 AI is analyzing the process to identify key moments..."):
            prompt = f"""
            You are an expert at analyzing process documentation and identifying key moments that should be captured in screenshots.
            
            Analyze this transcript of someone explaining an accounting or business process. The transcript includes timestamps.
            
            Identify timestamps where screenshots should be captured, specifically when the person describes:
            
            1. **Navigation paths** - When they describe how to get somewhere (e.g., "go to Menu > Options > Add Account", "click on the Transactions tab", "navigate to Reports")
            2. **Clicking buttons or links** - Any action that opens something new
            3. **Entering data** - When they're filling out forms or fields
            4. **Submitting or saving** - Final actions in a step
            5. **Key decision points** - Where choices need to be made
            6. **Opening menus or dropdowns** - Especially navigation menus
            7. **Important screens or pages** - When they arrive at a significant location in the system
            
            **IMPORTANT FOR NAVIGATION**: If someone describes a navigation path like "Menu > Options > Add Account", that should be ONE screenshot showing that final destination, not multiple screenshots.
            
            Return your analysis as a JSON array with this exact structure:
            [
                {{
                    "timestamp": "MM:SS",
                    "type": "navigation|action|data_entry|decision|submission",
                    "description": "Brief description of what's happening",
                    "navigation_path": "Menu > Options > Add Account" (only if type is navigation, otherwise null)
                }}
            ]
            
            Only include the JSON array in your response, no other text.
            
            TRANSCRIPT:
            {transcript}
            """
            
            response = model.generate_content(prompt)
            
            # Extract JSON from response
            response_text = response.text.strip()
            
            # Remove markdown code blocks if present
            if response_text.startswith('```json'):
                response_text = response_text[7:]
            if response_text.startswith('```'):
                response_text = response_text[3:]
            if response_text.endswith('```'):
                response_text = response_text[:-3]
            
            response_text = response_text.strip()
            
            # Parse JSON
            key_moments = json.loads(response_text)
            
            return key_moments
            
    except json.JSONDecodeError as e:
        st.error(f"Error parsing AI response: {str(e)}")
        st.error(f"AI Response was: {response_text}")
        return None
    except Exception as e:
        st.error(f"Error analyzing transcript: {str(e)}")
        import traceback
        st.error(f"Details: {traceback.format_exc()}")
        return None

def timestamp_to_seconds(timestamp_str):
    """Convert timestamp string (MM:SS or HH:MM:SS) to seconds"""
    try:
        parts = timestamp_str.strip().split(':')
        if len(parts) == 2:
            minutes, seconds = parts
            return int(minutes) * 60 + float(seconds)
        elif len(parts) == 3:
            hours, minutes, seconds = parts
            return int(hours) * 3600 + int(minutes) * 60 + float(seconds)
    except:
        return 0

def seconds_to_timestamp(seconds):
    """Convert seconds to timestamp string (MM:SS)"""
    mins = int(seconds // 60)
    secs = int(seconds % 60)
    return f"{mins}:{secs:02d}"

def extract_frame_at_timestamp(video_path, timestamp_seconds):
    """Extract a single frame from video at given timestamp"""
    try:
        cap = cv2.VideoCapture(video_path)
        
        # Set position to timestamp (in milliseconds)
        cap.set(cv2.CAP_PROP_POS_MSEC, timestamp_seconds * 1000)
        
        # Read frame
        success, frame = cap.read()
        cap.release()
        
        if success:
            # Convert BGR to RGB
            frame_rgb = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
            # Convert to PIL Image
            img = Image.fromarray(frame_rgb)
            return img
        else:
            return None
            
    except Exception as e:
        st.error(f"Error extracting frame at {timestamp_seconds}s: {str(e)}")
        return None

def extract_all_frames(video_path, key_moments):
    """Extract frames for all key moments"""
    frames = []
    
    progress_bar = st.progress(0)
    status_text = st.empty()
    
    for i, moment in enumerate(key_moments):
        timestamp_seconds = timestamp_to_seconds(moment['timestamp'])
        status_text.text(f"Extracting frame {i+1}/{len(key_moments)} at {moment['timestamp']}...")
        
        frame = extract_frame_at_timestamp(video_path, timestamp_seconds)
        
        if frame:
            frames.append({
                'moment': moment,
                'image': frame,
                'timestamp': moment['timestamp']
            })
        
        progress_bar.progress((i + 1) / len(key_moments))
    
    status_text.empty()
    progress_bar.empty()
    
    return frames

def image_to_base64(image):
    """Convert PIL Image to base64 string"""
    buffered = io.BytesIO()
    image.save(buffered, format="PNG")
    img_str = base64.b64encode(buffered.getvalue()).decode()
    return f"data:image/png;base64,{img_str}"

def generate_documentation(transcript, frames):
    """Generate final documentation with Gemini using delimiters"""
    try:
        with st.spinner("🤖 Generating professional documentation... This may take 1-2 minutes..."):
            content_parts = []
            
            instructions = """
You are an expert technical writer specializing in accounting and finance process documentation.

I will provide you with:
1. A transcript of someone explaining a process
2. Screenshots at key moments in the process

Please create professional, audit-ready Standard Operating Procedure (SOP) documentation.

**CRITICAL: Use ONLY these delimiters. Do NOT use markdown formatting (no **, ##, --, etc):**
- [TITLE] for the main process title
- [SECTION] for major section headers (Purpose, Scope, Prerequisites, Step-by-Step Instructions, Control Points, Common Issues & Troubleshooting, Frequency)
- [SUBSECTION] for subsection headers
- [STEP] for steps - Word will auto-number these - just provide the description without manual numbers
- [BULLET] for bullet points
- [SCREENSHOT] on its own line to indicate where a screenshot should be placed

**Document Structure:**
1. [TITLE] Process Title - Clear, descriptive title
2. [SECTION] Purpose - Why this process exists and what it accomplishes
3. [SECTION] Scope - What this process covers
4. [SECTION] Prerequisites - What needs to be in place before starting
5. [SECTION] Step-by-Step Instructions - Step descriptions (Word will auto-number as 1, 2, 3, etc.)
6. [SECTION] Control Points - Key moments where accuracy is critical (for SOX compliance)
7. [SECTION] Common Issues & Troubleshooting - Potential problems and solutions
8. [SECTION] Frequency - How often this process is performed

**Writing Guidelines:**
- Use imperative mood: "Click the Submit button" not "You click"
- Be concise but complete
- Use proper accounting terminology
- Make it audit-ready with clear accountability and verification steps
- For lists within sections, use [BULLET] instead of numbered lists
- Reference screenshots by placing [SCREENSHOT] on its own line where each should appear
- Do NOT manually number steps - just use [STEP] and describe what to do, Word will auto-number

**CRITICAL FORMATTING RULES:**
- EVERY line MUST start with a delimiter: [TITLE], [SECTION], [SUBSECTION], [STEP], [BULLET], or be a continuation line
- Do NOT use markdown: no **, ##, [[, ]], --, etc.
- Do NOT include manual numbers (1., 2., 3.) in [STEP] lines - let Word auto-number them
- No other formatting - just plain text

**Example Format:**
[TITLE] Account Reconciliation Process

[SECTION] Purpose
This process ensures all accounts match bank statements and are properly reconciled on a monthly basis.

[SECTION] Prerequisites
[BULLET] Active login credentials for the accounting system
[BULLET] Current month bank statements
[BULLET] Prior month reconciliation file

[SECTION] Step-by-Step Instructions
[STEP] Open the accounting system using your login credentials.
[SCREENSHOT]
[STEP] Navigate to Accounting menu, then select Reconciliation module.
[SCREENSHOT]
[STEP] Select the account to reconcile from the dropdown list.
[BULLET] Note: Contact IT if you don't have access to an account
[BULLET] You should see all active accounts in this list
[STEP] Upload the current bank statement CSV file.
[SCREENSHOT]

Notice: [STEP] has NO manual numbers - Word will automatically number these as 1, 2, 3, etc.
This allows Word to automatically renumber if steps are added or removed later.

Remember: Start each line with a delimiter. No manual numbers. No markdown. Keep formatting simple and clean.

TRANSCRIPT:
{transcript}

SCREENSHOTS TO INCLUDE:
"""
            
            content_parts.append(instructions)
            
            # Add screenshot context
            for i, frame_data in enumerate(frames, 1):
                moment = frame_data['moment']
                screenshot_info = f"Screenshot {i} [{moment['timestamp']}]: {moment['description']}"
                if moment.get('navigation_path'):
                    screenshot_info += f" | Navigation: {moment['navigation_path']}"
                content_parts.append(screenshot_info)
                content_parts.append(frame_data['image'])
            
            # Generate documentation
            response = model.generate_content(content_parts)
            
            return response.text
            
    except Exception as e:
        st.error(f"Error generating documentation: {str(e)}")
        import traceback
        st.error(f"Details: {traceback.format_exc()}")
        return None

def add_paragraph_with_screenshots(doc, text, frames, style=None):
    """Add a paragraph and embed screenshots referenced with [SCREENSHOT]"""
    text_stripped = text.strip()
    
    if not text_stripped:
        return
    
    # Add the paragraph
    if style:
        doc.add_paragraph(text_stripped, style=style)
    else:
        doc.add_paragraph(text_stripped)

def add_screenshot(doc, screenshot_num, frames):
    """Add a screenshot image to the document"""
    if screenshot_num <= len(frames):
        frame_data = frames[screenshot_num - 1]
        
        # Add image
        img_buffer = io.BytesIO()
        frame_data['image'].save(img_buffer, format='PNG')
        img_buffer.seek(0)
        doc.add_picture(img_buffer, width=Inches(6))
        
        # Add caption
        caption = doc.add_paragraph()
        caption_run = caption.add_run(f"Screenshot {screenshot_num}: {frame_data['moment']['description']}")
        caption_run.italic = True
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

def create_word_document(content, frames):
    """Create a Word document by parsing delimiter-based content"""
    try:
        doc = Document()
        
        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        
        lines = content.split('\n')
        i = 0
        screenshot_counter = 0
        
        while i < len(lines):
            line = lines[i].strip()
            
            # Skip empty lines
            if not line:
                i += 1
                continue
            
            # Handle TITLE
            if line.startswith('[TITLE]'):
                title_text = line[7:].strip()
                p = doc.add_heading(title_text, level=0)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Handle SECTION
            elif line.startswith('[SECTION]'):
                section_text = line[9:].strip()
                doc.add_heading(section_text, level=1)
            
            # Handle SUBSECTION
            elif line.startswith('[SUBSECTION]'):
                subsection_text = line[12:].strip()
                doc.add_heading(subsection_text, level=2)
            
            # Handle STEP
            elif line.startswith('[STEP]'):
                step_text = line[6:].strip()
                add_paragraph_with_screenshots(doc, step_text, frames, style='List Number')
            
            # Handle BULLET
            elif line.startswith('[BULLET]'):
                bullet_text = line[8:].strip()
                add_paragraph_with_screenshots(doc, bullet_text, frames, style='List Bullet')
            
            # Handle SCREENSHOT
            elif line.startswith('[SCREENSHOT]'):
                screenshot_counter += 1
                add_screenshot(doc, screenshot_counter, frames)
            
            # Regular text (continuation or standalone)
            else:
                if line:  # Only add non-empty lines
                    add_paragraph_with_screenshots(doc, line, frames)
            
            i += 1
        
        # Add footer
        section = doc.sections[0]
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.text = f"Generated by AI Process Documentation Generator on {datetime.now().strftime('%B %d, %Y')}"
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Save to bytes
        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        
        return doc_buffer.getvalue()
        
    except Exception as e:
        st.error(f"Error creating Word document: {str(e)}")
        import traceback
        st.error(f"Details: {traceback.format_exc()}")
        return None

def upload_word_doc_to_drive(word_doc_bytes, creds):
    """Upload an existing Word document to Google Drive"""
    try:
        from googleapiclient.discovery import build
        from googleapiclient.http import MediaIoBaseUpload
        
        # Build the Drive service
        drive_service = build('drive', 'v3', credentials=creds)
        
        # Prepare the file
        file_metadata = {
            'name': f"Process Documentation {datetime.now().strftime('%Y-%m-%d %H-%M')}.docx",
            'mimeType': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        }
        
        # Upload
        media = MediaIoBaseUpload(
            io.BytesIO(word_doc_bytes),
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            resumable=True
        )
        
        file = drive_service.files().create(
            body=file_metadata,
            media_body=media,
            fields='id,webViewLink'
        ).execute()
        
        file_id = file.get('id')
        file_url = file.get('webViewLink')
        
        return file_url, file_id
        
    except Exception as e:
        st.error(f"Error uploading to Google Drive: {str(e)}")
        import traceback
        st.error(f"Details: {traceback.format_exc()}")
        return None, None

def authenticate_google():
    """Authenticate with Google using a hosted Web OAuth flow (Streamlit Cloud compatible).

    Behavior:
    - If the current request contains an OAuth callback (code/state), exchange it for tokens and return creds.
    - Otherwise, generate an authorization URL and prompt the user to authenticate, then return None.
    """
    try:
        from google_auth_oauthlib.flow import Flow
        
        SCOPES = ['https://www.googleapis.com/auth/documents',
                  'https://www.googleapis.com/auth/drive.file']
        
        # Load OAuth settings from Streamlit secrets
        # Expected structure in .streamlit/secrets.toml:
        # [google_oauth]
        # client_id = "..."
        # client_secret = "..."
        # redirect_uri = "https://<your-streamlit-app-url>"
        try:
            oauth_cfg = st.secrets["google_oauth"]
            client_id = oauth_cfg["client_id"]
            client_secret = oauth_cfg["client_secret"]
            redirect_uri = oauth_cfg["redirect_uri"]
        except Exception:
            st.error("⚠️ Google OAuth is not configured. Add google_oauth.client_id, client_secret, and redirect_uri to Streamlit Secrets.")
            return None
        
        client_config = {
            "web": {
                "client_id": client_id,
                "project_id": "streamlit-sop",
                "auth_uri": "https://accounts.google.com/o/oauth2/auth",
                "token_uri": "https://oauth2.googleapis.com/token",
                "auth_provider_x509_cert_url": "https://www.googleapis.com/oauth2/v1/certs",
                "client_secret": client_secret,
                "redirect_uris": [redirect_uri]
            }
        }
        
        # Read query params compatibly across Streamlit versions
        try:
            query_params = st.query_params  # Streamlit >= 1.29
        except Exception:
            query_params = st.experimental_get_query_params()  # 1.28 fallback
        
        code = query_params.get("code") if isinstance(query_params, dict) else None
        if isinstance(code, list):
            code = code[0] if code else None
        state_param = query_params.get("state") if isinstance(query_params, dict) else None
        if isinstance(state_param, list):
            state_param = state_param[0] if state_param else None
        
        # If returning from Google's OAuth consent screen, finalize the flow
        if code:
            flow = Flow.from_client_config(client_config=client_config, scopes=SCOPES, redirect_uri=redirect_uri)
            # Optional CSRF protection if we previously stored state
            if state_param and "oauth_state" in st.session_state and st.session_state.oauth_state != state_param:
                st.error("OAuth state mismatch. Please try authenticating again.")
                return None
            
            # Exchange code for tokens
            flow.fetch_token(code=code)
            creds = flow.credentials
            
            # Clear query params to clean the URL (best-effort)
            try:
                st.experimental_set_query_params()  # Clears params in 1.28
            except Exception:
                pass
            
            return creds
        
        # Otherwise, initiate the OAuth flow by generating the authorization URL
        flow = Flow.from_client_config(client_config=client_config, scopes=SCOPES, redirect_uri=redirect_uri)
        authorization_url, state = flow.authorization_url(
            access_type='offline',
            include_granted_scopes='true',
            prompt='consent'
        )
        st.session_state.oauth_state = state
        
        # Present link for user to authenticate; after redirect back, the above code-path will run
        st.markdown(f"[Click here to authenticate with Google]({authorization_url})")
        st.info("After granting access, you will be redirected back here automatically.")
        return None
        
    except Exception as e:
        st.error(f"Authentication error: {str(e)}")
        import traceback
        st.error(f"Details: {traceback.format_exc()}")
        return None

def show_image_viewer(frames):
    """Display the full-screen image viewer with navigation"""
    if not frames or st.session_state.viewing_image is None:
        return
    
    current_index = st.session_state.viewing_image
    total_images = len(frames)
    
    # Ensure index is valid
    if current_index < 0 or current_index >= total_images:
        st.session_state.viewing_image = None
        return
    
    current_frame = frames[current_index]
    
    # Create the viewer container
    st.markdown("""
    <div class="image-viewer-container">
    """, unsafe_allow_html=True)
    
    st.markdown("### 🖼️ Image Viewer")
    
    # Navigation controls
    col1, col2, col3, col4, col5 = st.columns([1, 1, 3, 1, 1])
    
    with col1:
        if st.button("← Previous", disabled=(current_index == 0)):
            st.session_state.viewing_image = current_index - 1
            st.rerun()
    
    with col2:
        st.markdown(f"<div style='text-align: center; padding: 8px;'><strong>{current_index + 1} / {total_images}</strong></div>", unsafe_allow_html=True)
    
    with col4:
        if st.button("Next →", disabled=(current_index == total_images - 1)):
            st.session_state.viewing_image = current_index + 1
            st.rerun()
    
    with col5:
        if st.button("✕ Close Viewer", type="secondary"):
            st.session_state.viewing_image = None
            st.rerun()
    
    st.markdown("</div>", unsafe_allow_html=True)
    
    # Display full-size image - centered and slightly larger (550px)
    col_left, col_img, col_right = st.columns([1, 2, 1])
    with col_img:
        st.image(current_frame['image'], width=550)
    
    # Image details
    st.markdown("---")
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown(f"**⏱️ Timestamp:** {current_frame['timestamp']}")
        st.markdown(f"**📍 Type:** {current_frame['moment']['type']}")
    
    with col2:
        st.markdown(f"**📝 Description:**\n{current_frame['moment']['description']}")
        if current_frame['moment'].get('navigation_path'):
            st.markdown(f"**🧭 Navigation:** `{current_frame['moment']['navigation_path']}`")

def show_moment_editor(key_moments):
    """Show interactive editor for key moments"""
    st.markdown("### 📝 Edit Key Moments")
    st.info("Review and edit the AI-identified moments. Select moments to delete, or add new ones.")
    
    # Check if there are pending new moments to add
    if 'pending_new_moment' in st.session_state and st.session_state.pending_new_moment:
        # Add them to key_moments for display
        key_moments = key_moments + st.session_state.pending_new_moment
        st.info(f"✨ {len(st.session_state.pending_new_moment)} new moment(s) added - will be saved when you click 'Apply Changes'")
    
    # Multi-select for deletion
    st.markdown("#### 🗑️ Delete Unwanted Moments")
    moment_options = [
        f"{i+1}. [{m['timestamp']}] {m['type']}: {m['description'][:50]}..." 
        for i, m in enumerate(key_moments)
    ]
    
    moments_to_delete = st.multiselect(
        "Select moments to DELETE (they will be removed when you apply changes):",
        options=range(len(key_moments)),
        format_func=lambda x: moment_options[x],
        key="delete_moments"
    )
    
    if moments_to_delete:
        st.warning(f"⚠️ {len(moments_to_delete)} moment(s) will be deleted when you click 'Apply Changes'")
    
    st.markdown("---")
    st.markdown("#### ✏️ Edit Existing Moments")
    
    # Show editable fields for moments NOT marked for deletion
    edited_moments = []
    
    for i, moment in enumerate(key_moments):
        # Skip moments marked for deletion
        if i in moments_to_delete:
            continue
            
        with st.expander(f"✏️ Edit Moment {i+1}: [{moment['timestamp']}] {moment['type']}", expanded=False):
            col1, col2 = st.columns(2)
            
            with col1:
                new_timestamp = st.text_input(
                    "Timestamp (MM:SS)",
                    value=moment['timestamp'],
                    key=f"edit_time_{i}"
                )
                
                type_options = ['navigation', 'action', 'data_entry', 'decision', 'submission']
                current_type = moment.get('type', 'action')
                if current_type not in type_options:
                    current_type = 'action'
                
                new_type = st.selectbox(
                    "Type",
                    options=type_options,
                    index=type_options.index(current_type),
                    key=f"edit_type_{i}"
                )
            
            with col2:
                new_description = st.text_area(
                    "Description",
                    value=moment['description'],
                    height=100,
                    key=f"edit_desc_{i}"
                )
                
                if new_type == 'navigation':
                    new_nav_path = st.text_input(
                        "Navigation Path",
                        value=moment.get('navigation_path', ''),
                        key=f"edit_nav_{i}",
                        placeholder="Menu > Options > Add Account"
                    )
                else:
                    new_nav_path = None
            
            # Add edited moment to list
            edited_moments.append({
                'timestamp': new_timestamp,
                'type': new_type,
                'description': new_description,
                'navigation_path': new_nav_path
            })
    
    st.markdown("---")
    
    # Add new moment section
    with st.expander("➕ Add New Moment"):
        st.markdown("Add a new screenshot moment that AI may have missed:")
        
        col1, col2 = st.columns(2)
        
        with col1:
            new_timestamp = st.text_input(
                "Timestamp (MM:SS)",
                key="new_time",
                placeholder="2:30",
                help="Format: MM:SS (e.g., 2:30 for 2 minutes 30 seconds)"
            )
            new_type = st.selectbox(
                "Type",
                options=['navigation', 'action', 'data_entry', 'decision', 'submission'],
                key="new_type"
            )
        
        with col2:
            new_description = st.text_area(
                "Description",
                key="new_desc",
                height=100,
                placeholder="Describe what's happening at this moment..."
            )
            
            if new_type == 'navigation':
                new_nav_path = st.text_input(
                    "Navigation Path",
                    key="new_nav",
                    placeholder="Menu > Options > Add Account"
                )
            else:
                new_nav_path = None
        
        if st.button("➕ Add This Moment"):
            if new_timestamp and new_description:
                # Store in session state to persist across reruns
                if 'pending_new_moment' not in st.session_state:
                    st.session_state.pending_new_moment = []
                
                st.session_state.pending_new_moment.append({
                    'timestamp': new_timestamp,
                    'type': new_type,
                    'description': new_description,
                    'navigation_path': new_nav_path
                })
                st.success("✅ Moment added! Click 'Apply Changes & Extract Frames' below to save.")
            else:
                st.error("❌ Please enter both timestamp and description")
    
    # Summary
    st.markdown("---")
    original_count = len(key_moments)
    deleted_count = len(moments_to_delete)
    final_count = len(edited_moments)
    
    col1, col2, col3 = st.columns(3)
    with col1:
        st.metric("Original Moments", original_count)
    with col2:
        st.metric("Will Delete", deleted_count, delta=f"-{deleted_count}" if deleted_count > 0 else None)
    with col3:
        st.metric("Final Count", final_count, delta=final_count - original_count)
    
    return edited_moments

def main():
    initialize_session_state()
    
    # Header
    st.markdown("""
    <div class="main-header">
        <h1>📹 AI Process Documentation Generator</h1>
        <p>Upload a screen recording → AI generates professional documentation</p>
    </div>
    """, unsafe_allow_html=True)
    
    # API Key Setup
    if not model:
        st.error("""
        **⚠️ Setup Required**: Gemini API key not configured.
        
        1. Get your API key from [Google AI Studio](https://makersuite.google.com/app/apikey)
        2. Create a `.env` file with: `GEMINI_API_KEY=your_key_here`
        3. Restart the app
        """)
        return
    
    # Optional: Authenticate with Google first to avoid session loss on redirect
    st.markdown("---")
    st.markdown("### Step 0 (Optional, recommended): Authenticate with Google Drive")
    st.caption("Authenticate now to enable Drive upload later without losing progress during the redirect.")
    
    # If returning from Google OAuth, finalize and store credentials early
    try:
        _qp_top = st.query_params
    except Exception:
        _qp_top = st.experimental_get_query_params()
    if isinstance(_qp_top, dict) and _qp_top.get("code") and not st.session_state.get("oauth_processed"):
        _creds_top = authenticate_google()
        if _creds_top:
            st.session_state.google_creds = _creds_top
            st.session_state.oauth_processed = True
            st.success("✅ Google authentication complete.")
            try:
                st.experimental_set_query_params()
            except Exception:
                pass
            st.rerun()
    
    col_auth0a, col_auth0b = st.columns([1, 2])
    with col_auth0a:
        if st.button("🔐 Authenticate with Google now"):
            _ = authenticate_google()
    with col_auth0b:
        if st.session_state.get("google_creds"):
            st.success("Already authenticated for this session.")
        else:
            st.info("You can skip this and still download the Word file locally.")
    
    # Instructions
    with st.expander("ℹ️ How to Use This Tool", expanded=False):
        st.markdown("""
        ### Quick Start Guide
        
        1. **Record your process** using any screen recording tool (Zoom, Loom, QuickTime, OBS)
        2. **Talk through the process** as you demonstrate - explain navigation, actions, and decisions
        3. **Upload your video** (MP4, MOV, AVI, WebM)
        4. **AI analyzes** and identifies key moments
        5. **Review and edit** the identified moments (add, remove, or modify)
        6. **Extract frames** to see all screenshots
        7. **Generate documentation** - AI creates professional SOP
        8. **Download** your documentation
        
        ### Tips for Best Results
        
        - Speak clearly and describe what you're doing
        - Mention navigation paths explicitly
        - Keep recordings under 10 minutes
        - Explain the "why" not just the "what"
        """)
    
    # Main interface
    st.markdown("### Step 1: Upload Your Screen Recording")
    
    uploaded_file = st.file_uploader(
        "Choose a video file",
        type=['mp4', 'mov', 'avi', 'webm', 'mkv'],
        help="Upload a screen recording where you talk through a process"
    )
    
    if uploaded_file is not None:
        # Clear previous session artifacts ONLY when the uploaded file changes
        current_file_id = f"{uploaded_file.name}:{uploaded_file.size}"
        if st.session_state.get('last_uploaded_file_id') != current_file_id:
            st.session_state.transcript = None
            st.session_state.key_moments = None
            st.session_state.extracted_frames = None
            st.session_state.word_doc_bytes = None
            st.session_state.viewing_image = None
            st.session_state.audio_path = None
            st.session_state.video_path = None
            st.session_state.last_uploaded_file_id = current_file_id
        # Suppress extra file info rows; only show success once saved
        
        # Save uploaded file temporarily with proper flushing
        with tempfile.NamedTemporaryFile(delete=False, suffix='.mp4') as tmp_file:
            bytes_data = uploaded_file.getvalue()
            tmp_file.write(bytes_data)
            tmp_file.flush()
            os.fsync(tmp_file.fileno())
            video_path = tmp_file.name
            st.session_state.video_path = video_path
        
        # Verify file was saved correctly
        if os.path.exists(video_path) and os.path.getsize(video_path) > 0:
            st.success(f"✅ Video uploaded successfully!")
        else:
            st.error("❌ Video upload failed. Please try again.")
            video_path = None
        
        # Show video preview (only if file exists)
        if video_path:
            st.video(video_path)
            
            col1, col2 = st.columns([1, 3])
            with col1:
                process_button = st.button("🚀 Process Video", type="primary")
            
            if process_button:
                # Extract audio
                audio_path, duration = extract_audio_from_video(video_path)
                
                if audio_path:
                    st.session_state.audio_path = audio_path
                    st.success(f"✅ Audio extracted successfully! Video duration: {duration:.1f} seconds")
                    
                    # Transcribe audio
                    transcript = transcribe_audio_with_gemini(audio_path)
                    
                    if transcript:
                        st.session_state.transcript = transcript
                        st.success("✅ Transcription complete!")
                        
                        # Analyze transcript
                        key_moments = analyze_transcript_for_key_moments(transcript)
                        
                        if key_moments:
                            st.session_state.key_moments = key_moments
                            st.success(f"✅ Analysis complete! Found {len(key_moments)} key moments")
                            st.rerun()
    
    # Display transcript
    if st.session_state.transcript:
        st.markdown("---")
        st.markdown("### Step 2: Review Transcript")
        
        with st.expander("📝 Full Transcript", expanded=False):
            st.text_area(
                "Transcript",
                st.session_state.transcript,
                height=300,
                disabled=True
            )
    
    # Edit key moments
    if st.session_state.key_moments and not st.session_state.extracted_frames:
        st.markdown("---")
        st.markdown("### Step 3: Review & Edit Key Moments")
        
        st.info(f"🎯 AI identified **{len(st.session_state.key_moments)}** moments")
        
        # Show editor
        edited_moments = show_moment_editor(st.session_state.key_moments)
        
        # Apply changes button
        col1, col2, col3 = st.columns([2, 2, 2])
        
        with col1:
            if st.button("✅ Apply Changes & Extract Frames", type="primary"):
                # Clear pending new moments
                if 'pending_new_moment' in st.session_state:
                    del st.session_state.pending_new_moment
                
                # IMPORTANT: Use edited_moments which already has deletions applied
                if not edited_moments:
                    st.error("Cannot extract frames - no moments remaining after edits!")
                else:
                    # Sort moments by timestamp (chronological order)
                    sorted_moments = sorted(
                        edited_moments,
                        key=lambda m: timestamp_to_seconds(m['timestamp'])
                    )
                    
                    # Update moments with sorted version
                    st.session_state.key_moments = sorted_moments
                    
                    # Clear existing frames and documentation to force regeneration
                    st.session_state.extracted_frames = None
                    st.session_state.word_doc_bytes = None
                    
                    # Clear multiselect by removing its key from session state
                    if 'delete_moments' in st.session_state:
                        del st.session_state['delete_moments']
                    
                    # Extract frames in chronological order
                    with st.spinner("📸 Extracting screenshots from video..."):
                        frames = extract_all_frames(st.session_state.video_path, sorted_moments)
                        st.session_state.extracted_frames = frames
                    
                    # Delete video file after extraction to save storage
                    if os.path.exists(st.session_state.video_path):
                        os.remove(st.session_state.video_path)
                    
                    st.success(f"✅ Extracted {len(frames)} screenshots in chronological order!")
                    st.rerun()
        
        with col2:
            # Download edited moments as JSON
            json_data = json.dumps(edited_moments, indent=2)
            st.download_button(
                label="📥 Download as JSON",
                data=json_data,
                file_name=f"key_moments_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
                mime="application/json"
            )
    
    # Show image viewer if active
    if st.session_state.viewing_image is not None and st.session_state.extracted_frames:
        st.markdown("---")
        show_image_viewer(st.session_state.extracted_frames)
        st.markdown("---")
    
    # Show extracted frames
    if st.session_state.extracted_frames:
        st.markdown("---")
        st.markdown("### Step 4: Review Extracted Screenshots")
        
        st.success(f"✅ {len(st.session_state.extracted_frames)} screenshots ready")
        
        # Display frames in grid with click to view
        cols_per_row = 3
        for i in range(0, len(st.session_state.extracted_frames), cols_per_row):
            cols = st.columns(cols_per_row)
            for j in range(cols_per_row):
                if i + j < len(st.session_state.extracted_frames):
                    frame_data = st.session_state.extracted_frames[i + j]
                    frame_index = i + j
                    with cols[j]:
                        # Show thumbnail
                        st.image(frame_data['image'], caption=f"{frame_data['timestamp']} - {frame_data['moment']['type']}")
                        st.caption(frame_data['moment']['description'][:100] + "...")
                        
                        # View button
                        if st.button("🔍 View Full Size", key=f"view_{frame_index}"):
                            st.session_state.viewing_image = frame_index
                            st.rerun()
        
        # Generate documentation button
        st.markdown("---")
        
        if not st.session_state.word_doc_bytes:
            if st.button("🤖 Generate Professional Documentation", type="primary"):
                # Generate documentation using delimiters
                doc_content = generate_documentation(
                    st.session_state.transcript,
                    st.session_state.extracted_frames
                )
                
                if doc_content:
                    # Immediately create Word document
                    with st.spinner("📝 Creating Word document with embedded screenshots..."):
                        word_doc = create_word_document(
                            doc_content,
                            st.session_state.extracted_frames
                        )
                    
                    if word_doc:
                        st.session_state.word_doc_bytes = word_doc
                        st.success("✅ Documentation generated and ready to download!")
                        st.rerun()
        else:
            st.success("✅ Documentation ready to download!")
    
    # Show download options
    if st.session_state.word_doc_bytes:
        st.markdown("---")
        st.markdown("### Step 5: Download Your Professional Documentation")
        
        # Download Word document (single column)
        st.download_button(
            label="📥 Download Word Document (.docx)",
            data=st.session_state.word_doc_bytes,
            file_name=f"process_documentation_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            type="primary"
        )
        st.caption("Professional SOP with embedded screenshots")
        
        # Google Drive upload section
        st.markdown("---")
        st.markdown("### ☁️ Upload to Google Drive (Optional)")
        
        if not st.session_state.google_creds:
            # If returning from Google OAuth, finalize and store credentials
            try:
                _qp = st.query_params
            except Exception:
                _qp = st.experimental_get_query_params()
            if isinstance(_qp, dict) and _qp.get("code"):
                _creds = authenticate_google()
                if _creds:
                    st.session_state.google_creds = _creds
                    st.success("✅ Authentication successful!")
                    st.rerun()

            st.info("""
            💡 **Want to save this to your Google Drive?**
            
            Authenticate once, then you can upload the Word document directly to your Drive.
            
            **One-time setup required:**
            - You need `credentials.json` from Google Cloud Console
            - See instructions in sidebar or README
            """)
            
            col_auth1, col_auth2 = st.columns([1, 2])
            with col_auth1:
                if st.button("🔐 Authenticate with Google"):
                    # Initiates the web OAuth flow (opens consent screen). After redirect back, credentials are finalized automatically above.
                    _ = authenticate_google()
            
            with col_auth2:
                with st.expander("📖 Google Drive Setup Instructions"):
                    st.markdown("""
                    1. Go to [Google Cloud Console](https://console.cloud.google.com/)
                    2. Create a project
                    3. Enable Google Drive API
                    4. Create OAuth 2.0 Desktop credentials
                    5. Download as `credentials.json`
                    6. Place in app directory
                    """)
        else:
            st.success("✅ Authenticated with Google")
            
            col_upload1, col_upload2 = st.columns([1, 1])
            
            with col_upload1:
                if st.button("☁️ Upload to Google Drive", type="primary"):
                    with st.spinner("Uploading to Google Drive..."):
                        file_url, file_id = upload_word_doc_to_drive(
                            st.session_state.word_doc_bytes,
                            st.session_state.google_creds
                        )
                    
                    if file_url:
                        st.success("✅ Uploaded successfully!")
                        st.markdown(f"**[Open in Google Drive]({file_url})**")
                        st.info("💡 Tip: Click 'Open with Google Docs' in Drive to convert to Google Doc format if desired.")
                    else:
                        st.error("Upload failed. Please try downloading instead.")
            
            # Removed 'Change Google Account' to reduce confusion for demo
    
    # Sidebar
    with st.sidebar:
        if st.button("🔄 Start New Process"):
            for key in list(st.session_state.keys()):
                del st.session_state[key]
            st.rerun()
        
        st.markdown("---")
        
        # Tips for Best Results
        with st.expander("💡 Tips for Best Results", expanded=False):
            st.markdown("""
            **Video Recording:**
            - Use a clear screen recording tool
            - Speak clearly and descriptively
            - Mention navigation paths explicitly (e.g., "Menu > Options > Add Account")
            
            **Pro Tip for Better Results:**
            - In your video, say "**take a screenshot here**" at key moments
            - Describe what should be highlighted (e.g., "the Export button")
            - This helps AI sync audio to visuals more accurately
            
            **General Tips:**
            - Keep recordings under 10 minutes
            - Explain the "why" not just the "what"
            - Review and edit moments before extracting frames
            """)
        
        # Roadmap
        with st.expander("🗺️ Roadmap & Future Features", expanded=False):
            st.markdown("""
            **Phase 2 - Smart Annotations:**
            - Auto-detect and highlight UI elements referenced in audio
            - Intelligently annotate buttons, fields, and important areas
            - Include in final documentation for clarity
            
            **Phase 3 - Enhanced Frame Selection:**
            - Preview multiple frames around each timestamp
            - User-selected optimal screenshot capture
            - Better sync between audio and visuals
            
            **Phase 4 - Advanced Features:**
            - Multi-language support
            - Long-form video segmentation
            - Batch processing multiple videos
            - Integration with document management systems
            """)
        
        # Known Limitations
        with st.expander("⚠️ Known Limitations", expanded=False):
            st.markdown("""
            **Screenshot Timing:**
            - Screenshots may lag 1-2 seconds behind audio description
            - Solution: Say "take a screenshot here" explicitly in your video
            
            **AI Accuracy:**
            - Transcription works best with clear audio
            - Key moment detection depends on explicit descriptions
            
            **Video Requirements:**
            - Best results with videos under 10 minutes
            - Supports: MP4, MOV, AVI, WebM, MKV
            - Requires audio track for transcription
            
            **Current Scope:**
            - Optimized for single-process demonstrations
            - Best for standard business software workflows
            """)
        
        # Data Security
        with st.expander("🔒 Data Security", expanded=False):
            st.markdown("""
            **Video Storage:**
            - Videos for this prototype are stored as temporary files on Streamlit Cloud's servers
            - Automatically deleted after screenshot frames are extracted
            - Not retained on Streamlit Cloud's servers after processing completes
            - Enterprise version would store temp files on company private servers
            
            **API Processing:**
            - Audio and screenshots are sent to Google's Gemini-2.5 Pro API for analysis
            - In enterprise deployments, can be configured to use company's firewalled enterprise Gemini instance
            
            **Google Drive Integration:**
            - Generated documents are saved only to user's Google Drive
            - User explicitly authorizes the app to access their Drive
            - User retains full control and can delete files anytime
            - App cannot access other Drive files
            
            **Deployment:**
            - Prototype: Hosted on Streamlit Cloud
            - Enterprise: Deployed on private company servers
            - Recommended: Run on secured internal infrastructure for sensitive accounting data
            """)
        
        # Tools Used
        with st.expander("🛠️ Tools & Technologies", expanded=False):
            st.markdown("""
            **Core Framework:**
            - Streamlit 1.28+
            - Python 3.8+
            
            **AI**
            - Google Gemini-2.5 Pro API
            
            **Video & Image Processing:**
            - MoviePy 1.0+
            - OpenCV (cv2) 4.8+
            - Pillow (PIL) 10.0+
            
            **Document Generation:**
            - python-docx 0.8+
            
            **Google Integration:**
            - google-auth-oauthlib 1.1+
            - google-auth-httplib2 0.2+
            - google-api-python-client 2.100+
            
            **Environment & Configuration:**
            - python-dotenv 1.0+
            - pickle (built-in)
            
            **Deployment:**
            - Streamlit Cloud (prototype)
            - Self-hosted deployment (enterprise)
            """)
        
        # Accounting Use Cases
        with st.expander("📊 Accounting Use Cases", expanded=False):
            st.markdown("""
            - Training new team members
            - Sharing process information among team and stakeholders
            - Easy to update when processes change
            - Walkthrough support for auditors
            
            """)
        
        st.markdown("---")
        st.markdown("""
        ### 📝 About
        
        **AI Process Documentation Generator**
        
        Automatically converts screen recordings into professional Standard Operating Procedures (SOPs) using AI transcription and visual analysis.
        """)


if __name__ == "__main__":
    main()